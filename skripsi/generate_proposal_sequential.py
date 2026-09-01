# ============================================================================
# skripsi/generate_proposal_sequential.py
#
# Membangun proposal tugas akhir (Bab I-IV) untuk topik:
#   "Pencarian Model Pareto-Optimal Akurasi, Latensi, dan Ukuran Model
#    pada Pengenalan Karakter Optik Sekuensial Aksara Sunda"
#
# Format mengikuti Pedoman Penyusunan dan Penulisan Skripsi Unpad
# (Keputusan Rektor No. 391/H6.1/KEP/PP/2011) dan slide MP-05 Prodi TI:
#   - A4, margin atas/kiri 4 cm, bawah/kanan 3 cm
#   - Times New Roman 12, spasi ganda pada tubuh teks
#   - Indentasi baris pertama 1,25 cm (5 ketukan)
#   - Sitasi gaya Harvard (Author-Year)
#
# Jalankan: python generate_proposal_sequential.py
# ============================================================================
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Proposal_SekuensialAksaraSunda_Bim_Yusuf_Karang.docx")

JUDUL = [
    "PENCARIAN MODEL PARETO-OPTIMAL AKURASI, LATENSI,",
    "DAN UKURAN MODEL PADA PENGENALAN KARAKTER OPTIK",
    "SEKUENSIAL AKSARA SUNDA",
]


# ---------------------------------------------------------------- utilitas
def setup(doc):
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin, s.left_margin = Cm(4.0), Cm(4.0)
    s.bottom_margin, s.right_margin = Cm(3.0), Cm(3.0)
    n = doc.styles["Normal"]
    n.font.name, n.font.size = "Times New Roman", Pt(12)
    pf = n.paragraph_format
    pf.space_before, pf.space_after = Pt(0), Pt(0)
    pf.line_spacing = 2.0


def body(doc, text, indent=True, spacing=2.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = spacing
    if indent:
        pf.first_line_indent = Cm(1.25)
    p.add_run(text)
    return p


def center(doc, text, bold=False, spacing=1.5, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    p.add_run(text).bold = bold
    return p


def heading(doc, number, title, before=12, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    p.add_run(f"{number}\t{title}").bold = True
    return p


def bab(doc, nomor, judul):
    doc.add_page_break()
    center(doc, f"BAB {nomor}", bold=True, spacing=1.5)
    center(doc, judul, bold=True, spacing=1.5, after=12)


def listitem(doc, text, spacing=2.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.left_indent, pf.first_line_indent = Cm(1.25), Cm(-0.75)
    p.add_run(text)
    return p


def caption(doc, text, before=6, after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    p.add_run(text)
    return p


def table(doc, rows, widths=None, font_pt=11):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(font_pt)
            if i == 0:
                r.bold = True
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


def pustaka(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(12)
    pf.left_indent, pf.first_line_indent = Cm(1.0), Cm(-1.0)
    p.add_run(text)
    return p


# ============================================================ HALAMAN JUDUL
def halaman_judul(doc):
    for _ in range(2):
        center(doc, "", spacing=1.5)
    for baris in JUDUL:
        center(doc, baris, bold=True, spacing=1.5)
    center(doc, "", spacing=1.5)
    center(doc, "", spacing=1.5)
    center(doc, "PROPOSAL TUGAS AKHIR", bold=True, spacing=1.5)
    center(doc, "", spacing=1.5)
    center(doc, "diajukan untuk menempuh ujian sarjana", spacing=1.5)
    center(doc, "pada Program Studi S-1 Teknik Informatika", spacing=1.5)
    center(doc, "Universitas Padjadjaran", spacing=1.5)
    for _ in range(2):
        center(doc, "", spacing=1.5)
    center(doc, "Oleh", spacing=1.5)
    center(doc, "", spacing=1.5)
    center(doc, "BIM YUSUF KARANG", bold=True, spacing=1.5)
    center(doc, "NPM. 140810230084", spacing=1.5)
    for _ in range(3):
        center(doc, "", spacing=1.5)
    center(doc, "[LOGO UNIVERSITAS PADJADJARAN]", spacing=1.5)
    for _ in range(3):
        center(doc, "", spacing=1.5)
    center(doc, "PROGRAM STUDI S-1 TEKNIK INFORMATIKA", bold=True, spacing=1.5)
    center(doc, "DEPARTEMEN ILMU KOMPUTER", bold=True, spacing=1.5)
    center(doc, "FAKULTAS MATEMATIKA DAN ILMU PENGETAHUAN ALAM", bold=True, spacing=1.5)
    center(doc, "UNIVERSITAS PADJADJARAN", bold=True, spacing=1.5)
    center(doc, "JATINANGOR", bold=True, spacing=1.5)
    center(doc, "2026", bold=True, spacing=1.5)


# =========================================================== KATA PENGANTAR
def kata_pengantar(doc):
    doc.add_page_break()
    center(doc, "KATA PENGANTAR", bold=True, spacing=1.5, after=12)
    body(doc, "Puji dan syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa atas rahmat "
              "dan karunia-Nya sehingga proposal tugas akhir ini dapat diselesaikan. Proposal "
              "ini disusun sebagai salah satu syarat untuk menempuh ujian sarjana pada Program "
              "Studi S-1 Teknik Informatika, Departemen Ilmu Komputer, Fakultas Matematika dan "
              "Ilmu Pengetahuan Alam, Universitas Padjadjaran.")
    body(doc, "Penelitian yang diusulkan berangkat dari satu kesenjangan yang terdokumentasi "
              "secara eksplisit pada literatur mutakhir, yaitu belum tersedianya sistem "
              "pengenalan karakter optik sekuensial yang layak pakai untuk Aksara Sunda. "
              "Penelitian ini berupaya menjawab kesenjangan tersebut sekaligus mengkarakterisasi "
              "hubungan pertukaran antara akurasi, latensi, dan ukuran model sehingga hasilnya "
              "tidak hanya akurat secara akademis, tetapi juga dapat diterapkan pada perangkat "
              "dengan sumber daya komputasi terbatas.")
    body(doc, "Penulis mengucapkan terima kasih kepada dosen pembimbing atas arahan dan koreksi "
              "yang diberikan, kepada seluruh staf pengajar Program Studi S-1 Teknik Informatika "
              "atas ilmu yang telah dibagikan, serta kepada keluarga dan rekan-rekan atas "
              "dukungannya. Penulis menyadari proposal ini masih memiliki kekurangan, sehingga "
              "kritik dan saran yang membangun sangat penulis harapkan.")
    for _ in range(2):
        body(doc, "", indent=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    p.add_run("Jatinangor,          2026")
    for _ in range(2):
        body(doc, "", indent=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    p.add_run("Bim Yusuf Karang")


# ================================================================ DAFTAR ISI
DAFTAR_ISI = [
    "BAB I PENDAHULUAN",
    "1.1 Latar Belakang Penelitian",
    "1.2 Identifikasi Masalah",
    "1.3 Maksud dan Tujuan Penelitian",
    "1.4 Kegunaan Penelitian",
    "1.5 Kerangka Pemikiran",
    "1.6 Metodologi Penelitian",
    "1.7 Waktu dan Lokasi Penelitian",
    "BAB II TINJAUAN PUSTAKA",
    "2.1 Aksara Sunda dan Upaya Pelestarian Digital",
    "2.2 Pengenalan Karakter Optik: Dari Terisolasi ke Sekuensial",
    "2.3 Convolutional Neural Network sebagai Ekstraktor Fitur",
    "2.4 Connectionist Temporal Classification",
    "2.5 Paradigma Dekoder Alternatif: Attention dan Transformer",
    "2.6 Sintesis Data untuk Pelatihan Model Pengenalan Teks",
    "2.7 Efisiensi Model dan Optimalitas Pareto",
    "2.8 Penelitian Terdahulu yang Relevan",
    "2.9 Kesenjangan Penelitian",
    "BAB III METODE PENELITIAN",
    "3.1 Asas Penelitian",
    "3.2 Rancangan Penelitian",
    "3.3 Sumber dan Rancangan Data",
    "3.3.1 Sumber Data",
    "3.3.2 Prosedur Sintesis Citra Teks Sekuensial",
    "3.3.3 Skema Pembagian Data",
    "3.4 Arsitektur Model yang Diusulkan",
    "3.4.1 Tulang Punggung Konvolusi",
    "3.4.2 Kepala Sekuensial Berbasis CTC",
    "3.4.3 Paradigma Dekoder Pembanding",
    "3.5 Model Garis Dasar Segmentasi-lalu-Klasifikasi",
    "3.6 Rancangan Eksperimen Dua Faktor",
    "3.7 Prosedur Pelatihan dan Pengolahan Data Eksperimen",
    "3.8 Pengukuran dan Kriteria Keberhasilan Penelitian",
    "3.9 Kebutuhan Perangkat Keras dan Perangkat Lunak",
    "3.10 Batasan Penelitian",
    "BAB IV JADWAL PENELITIAN",
    "4.1 Tahapan Pelaksanaan Penelitian",
    "4.2 Jadwal Kegiatan",
    "DAFTAR PUSTAKA",
    "RIWAYAT HIDUP",
]

DAFTAR_GAMBAR = [
    "Gambar 2.1  Pemetaan banyak-ke-satu pada Connectionist Temporal Classification",
    "Gambar 2.2  Ilustrasi frontier Pareto pada ruang akurasi-latensi-ukuran model",
    "Gambar 3.1  Diagram alir pelaksanaan penelitian",
    "Gambar 3.2  Prosedur sintesis citra kata Aksara Sunda dari korpus karakter terisolasi",
    "Gambar 3.3  Arsitektur CRNN yang diusulkan beserta dimensi peta fitur",
]

DAFTAR_TABEL = [
    "Tabel 3.1  Spesifikasi lapisan arsitektur CRNN yang diusulkan",
    "Tabel 3.2  Konfigurasi pengali lebar kanal pada rancangan dua faktor",
    "Tabel 3.3  Model pembanding dan garis dasar",
    "Tabel 3.4  Kriteria keberhasilan penelitian",
    "Tabel 4.1  Jadwal pelaksanaan penelitian",
]


def daftar(doc, judul, isi):
    doc.add_page_break()
    center(doc, judul, bold=True, spacing=1.5, after=12)
    for item in isi:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        if item.startswith("BAB") or item in ("DAFTAR PUSTAKA", "RIWAYAT HIDUP"):
            p.add_run(item).bold = True
        else:
            depth = item.count(".", 0, 6)
            pf.left_indent = Cm(0.75 * max(0, depth - 1))
            p.add_run(item)


# ==================================================================== BAB I
def bab_1(doc):
    bab(doc, "I", "PENDAHULUAN")

    heading(doc, "1.1", "Latar Belakang Penelitian", before=0)
    body(doc, "Aksara Sunda merupakan sistem tulisan tradisional masyarakat Sunda yang "
              "penggunaannya terdokumentasi pada prasasti dan naskah sejak abad ke-14. "
              "Penggunaannya menurun tajam pada masa kolonial seiring meluasnya aksara Latin, "
              "Pegon Arab, dan Cacarakan (Agustiansyah and Fauzi, 2025). Upaya revitalisasi "
              "dilakukan melalui kebijakan daerah dan pengintegrasian aksara ini ke dalam "
              "kurikulum muatan lokal di Jawa Barat dan Banten. Namun, keberlangsungan aksara "
              "daerah pada era digital tidak cukup ditopang oleh kebijakan pendidikan semata; "
              "diperlukan pula dukungan teknologi yang memungkinkan naskah dan teks beraksara "
              "Sunda diproses secara otomatis oleh mesin.")
    body(doc, "Teknologi yang menjembatani kebutuhan tersebut adalah pengenalan karakter optik "
              "atau Optical Character Recognition (OCR), yaitu konversi citra teks menjadi teks "
              "terbaca mesin. Penelitian OCR untuk Aksara Sunda yang tersedia hingga saat ini "
              "terkonsentrasi pada satu ranah sempit, yaitu klasifikasi karakter terisolasi. "
              "Agustiansyah and Fauzi (2025) membandingkan lima arsitektur pembelajaran mendalam "
              "(ResNet-50, MobileNetV2, EfficientNet-B0, ViT, dan DeiT) pada dataset seimbang "
              "berisi 30 kelas Aksara Sunda dan melaporkan akurasi hingga 96,9 persen pada data "
              "sebaran dalam. Capaian tersebut menunjukkan bahwa persoalan pengenalan satu "
              "karakter tunggal yang telah terpotong rapi secara praktis telah terselesaikan.")
    body(doc, "Persoalannya, teks Aksara Sunda di dunia nyata tidak hadir sebagai karakter "
              "tunggal yang terisolasi, melainkan sebagai rangkaian karakter yang membentuk kata "
              "dan kalimat. Kemampuan mengenali karakter tunggal tidak otomatis berarti kemampuan "
              "membaca teks utuh, karena pengenalan teks sekuensial menuntut model menyelaraskan "
              "sendiri antara wilayah citra dan urutan karakter tanpa anotasi posisi per karakter. "
              "Persoalan penyelarasan inilah yang membedakan pengenalan sekuensial dari "
              "klasifikasi terisolasi secara mendasar.")
    body(doc, "Bukti bahwa kesenjangan tersebut belum terjembatani disajikan secara eksplisit "
              "oleh Adilazuarda et al. (2025) melalui NusaAksara, tolok ukur multimodal dan "
              "multibahasa yang mencakup delapan aksara Nusantara, termasuk Aksara Sunda. Pada "
              "tugas OCR untuk Aksara Sunda, seluruh sistem yang diuji menghasilkan Character "
              "Error Rate (CER) melebihi satu, yang berarti jumlah kesalahan melampaui panjang "
              "teks acuan sehingga keluaran model tidak dapat digunakan sama sekali. Kegagalan "
              "tersebut tidak terbatas pada model bahasa visual berskala besar seperti GPT-4o "
              "dan Gemini Flash, tetapi juga terjadi pada PP-OCRv3 yang telah disetel khusus "
              "(fine-tuned) pada data tersebut. Penulis NusaAksara mengidentifikasi penyebab "
              "utamanya, yaitu keterbatasan data pelatihan yang sangat parah sehingga tidak "
              "memadai untuk pembelajaran yang efektif.")
    body(doc, "Diagnosis tersebut mengarahkan penelitian ini pada satu strategi kunci: apabila "
              "hambatan utamanya adalah ketersediaan data dan bukan ketiadaan algoritma yang "
              "sesuai, maka hambatan tersebut dapat ditembus melalui sintesis data. Strategi ini "
              "memiliki preseden yang kuat dan telah teruji pada ranah pengenalan teks. "
              "Jaderberg et al. (2014a) melatih pengenal teks pemandangan alam sepenuhnya dari "
              "data sintetis tanpa satu pun anotasi manusia dan tetap memperoleh performa "
              "terdepan pada masanya, sementara Gupta et al. (2016) menunjukkan hal serupa untuk "
              "tugas lokalisasi teks. Dalam konteks Aksara Sunda, korpus karakter terisolasi "
              "yang telah tersedia dapat dirangkai secara terprogram menjadi citra kata dan "
              "frasa, dengan label kebenaran yang bersifat eksak menurut konstruksinya. "
              "Pendekatan ini sekaligus menghilangkan kebutuhan akan tahap deteksi teks, karena "
              "citra yang disintesis telah terpotong pada tataran kata.")
    body(doc, "Kelayakan penerapan sebuah model OCR tidak ditentukan oleh akurasi semata. "
              "Penerapan pada aplikasi edukasi bergerak maupun perangkat digitalisasi lapangan "
              "menuntut model yang ringan dan cepat. Persoalannya, jumlah parameter tidak dapat "
              "diandalkan sebagai proksi bagi kecepatan sebenarnya. Ma et al. (2018) menunjukkan "
              "bahwa dua model dengan jumlah operasi setara dapat memiliki latensi nyata yang "
              "sangat berbeda akibat biaya akses memori dan tingkat paralelisme, sedangkan "
              "Dehghani et al. (2022) memperingatkan bahwa pengklaiman efisiensi berdasarkan satu "
              "metrik tunggal dapat menyesatkan karena metrik-metrik tersebut kerap saling "
              "bertentangan. Dengan demikian, pertanyaan yang tepat bukanlah model mana yang "
              "paling akurat, melainkan konfigurasi model mana yang optimal secara Pareto pada "
              "ruang tiga dimensi yang dibentuk oleh akurasi, latensi terukur, dan ukuran model.")
    body(doc, "Berdasarkan uraian tersebut, penelitian ini mengusulkan pembangunan sistem "
              "pengenalan karakter optik sekuensial untuk Aksara Sunda dengan memanfaatkan data "
              "sintetis, sekaligus mengkarakterisasi frontier Pareto pada ketiga sumbu evaluasi "
              "melalui rancangan eksperimen dua faktor yang memisahkan pengaruh ukuran model dari "
              "pengaruh rancangan arsitektur. Sejauh penelusuran penulis, karakterisasi semacam "
              "ini belum pernah dilakukan untuk pengenalan teks sekuensial Aksara Sunda.")

    heading(doc, "1.2", "Identifikasi Masalah")
    body(doc, "Berdasarkan uraian latar belakang di atas, permasalahan penelitian ini "
              "diidentifikasi sebagai berikut.")
    for t in [
        "1. Bagaimana merancang prosedur sintesis data teks sekuensial Aksara Sunda dari korpus "
        "karakter terisolasi yang tersedia sehingga diperoleh volume dan keragaman data yang "
        "memadai untuk melatih model pengenalan sekuensial?",
        "2. Bagaimana performa model pengenalan sekuensial berbasis Connectionist Temporal "
        "Classification pada citra teks Aksara Sunda hasil sintesis, diukur menggunakan "
        "Character Error Rate dan Word Error Rate?",
        "3. Bagaimana perbandingan performa model sekuensial terhadap pendekatan garis dasar "
        "berupa segmentasi karakter yang diikuti klasifikasi terisolasi?",
        "4. Bagaimana pengaruh penskalaan lebar kanal secara sistematis terhadap akurasi, "
        "latensi terukur, dan ukuran model, apabila topologi arsitektur dikendalikan konstan?",
        "5. Bagaimana perbandingan tiga paradigma dekoder sekuens, yaitu Connectionist Temporal "
        "Classification, attention, dan transformer, pada ketiga sumbu evaluasi tersebut?",
        "6. Konfigurasi model manakah yang menempati frontier Pareto pada ruang tiga dimensi "
        "akurasi, latensi, dan ukuran model untuk tugas pengenalan teks sekuensial Aksara Sunda?",
    ]:
        listitem(doc, t)

    heading(doc, "1.3", "Maksud dan Tujuan Penelitian")
    body(doc, "Penelitian ini dimaksudkan untuk merancang, mengimplementasikan, dan "
              "mengevaluasi secara komprehensif sistem pengenalan karakter optik sekuensial "
              "untuk Aksara Sunda, serta mengkarakterisasi hubungan pertukaran antara akurasi "
              "pengenalan, latensi inferensi, dan ukuran model pada sistem tersebut.")
    body(doc, "Secara khusus, tujuan penelitian ini adalah sebagai berikut. Pertama, "
              "merancang dan mengimplementasikan prosedur sintesis data teks sekuensial Aksara "
              "Sunda beserta korpus data yang dihasilkannya. Kedua, mengimplementasikan model "
              "pengenalan sekuensial berbasis Connectionist Temporal Classification dan mengukur "
              "performanya menggunakan Character Error Rate dan Word Error Rate. Ketiga, "
              "membandingkan model tersebut terhadap garis dasar segmentasi-lalu-klasifikasi "
              "serta terhadap paradigma dekoder attention dan transformer. Keempat, melaksanakan "
              "rancangan eksperimen dua faktor yang memisahkan pengaruh ukuran model dari "
              "pengaruh rancangan arsitektur. Kelima, mengidentifikasi himpunan konfigurasi yang "
              "optimal secara Pareto pada ruang akurasi, latensi, dan ukuran model.")

    heading(doc, "1.4", "Kegunaan Penelitian")
    heading(doc, "1.4.1", "Kegunaan Teoretis", before=6, after=4)
    body(doc, "Penelitian ini diharapkan memberikan kontribusi berupa dokumentasi empiris "
              "mengenai perilaku model pengenalan sekuensial pada aksara daerah dengan sumber "
              "daya data yang terbatas, suatu ranah yang menurut Adilazuarda et al. (2025) masih "
              "belum tergarap. Penelitian ini juga menguji apakah temuan mengenai ketidakandalan "
              "jumlah parameter sebagai proksi latensi, sebagaimana dilaporkan Ma et al. (2018) "
              "dan Dehghani et al. (2022) pada ranah klasifikasi citra, turut berlaku pada ranah "
              "pengenalan teks sekuensial resolusi rendah.")
    body(doc, "Selain itu, penerapan rancangan eksperimen dua faktor yang memisahkan pengaruh "
              "penskalaan lebar kanal dari pengaruh pemilihan arsitektur diharapkan memberikan "
              "landasan metodologis yang lebih tegas dibandingkan perbandingan antararsitektur "
              "pada ukuran bawaan masing-masing, yang lazim dijumpai pada literatur dan "
              "berpotensi mengaburkan kedua pengaruh tersebut.")
    heading(doc, "1.4.2", "Kegunaan Praktis", before=6, after=4)
    body(doc, "Secara praktis, penelitian ini diharapkan menghasilkan sistem pengenalan teks "
              "Aksara Sunda yang dapat menjadi fondasi bagi aplikasi pembelajaran aksara daerah, "
              "alat bantu transliterasi, maupun tahap awal digitalisasi naskah. Ketersediaan "
              "profil latensi terukur pada perangkat dengan sumber daya terbatas memungkinkan "
              "pengembang memilih konfigurasi model yang sesuai dengan anggaran komputasi "
              "perangkat sasaran, alih-alih memilih berdasarkan akurasi tertinggi semata.")
    body(doc, "Korpus data sintetis beserta perkakas sintesisnya juga direncanakan untuk "
              "didokumentasikan sehingga dapat digunakan kembali oleh peneliti lain yang "
              "menghadapi kendala serupa pada aksara daerah lainnya.")

    heading(doc, "1.5", "Kerangka Pemikiran")
    body(doc, "Kerangka pemikiran penelitian ini disusun secara deduktif dengan bertolak dari "
              "dua dalil yang telah mapan. Dalil pertama menyatakan bahwa pelabelan sekuens tanpa "
              "segmentasi eksplisit dapat dipelajari secara ujung-ke-ujung apabila fungsi "
              "kerugian menjumlahkan probabilitas atas seluruh penyelarasan yang sah antara "
              "keluaran per langkah waktu dan label acuan, sebagaimana dirumuskan Graves et al. "
              "(2006). Dalil kedua menyatakan bahwa data pelatihan sintetis dapat menggantikan "
              "data beranotasi manusia pada tugas pengenalan teks apabila proses sintesisnya "
              "menghadirkan keragaman yang memadai, sebagaimana ditunjukkan Jaderberg et al. "
              "(2014a) dan Gupta et al. (2016).")
    body(doc, "Dari kedua dalil tersebut dirumuskan hipotesis operasional bahwa model "
              "sekuensial berbasis Connectionist Temporal Classification yang dilatih pada citra "
              "teks Aksara Sunda hasil sintesis mampu mencapai Character Error Rate yang jauh di "
              "bawah satu, yaitu suatu tingkat yang seluruh sistem terdahulu pada tolok ukur "
              "NusaAksara gagal mencapainya. Hipotesis kedua menyatakan bahwa hubungan antara "
              "ukuran model dan akurasi pada ranah ini tidak bersifat linear, sehingga terdapat "
              "konfigurasi berukuran lebih kecil yang mendominasi konfigurasi berukuran lebih "
              "besar pada ketiga sumbu evaluasi secara bersamaan.")
    body(doc, "Pengujian kedua hipotesis dilakukan melalui rancangan eksperimen dua faktor. "
              "Faktor pertama adalah rancangan arsitektur, yang dikendalikan konstan dalam satu "
              "kurva dan divariasikan antarkurva. Faktor kedua adalah pengali lebar kanal, yang "
              "divariasikan secara sistematis di dalam satu arsitektur yang sama. Bentuk setiap "
              "kurva merepresentasikan pengaruh murni ukuran model, sedangkan jarak vertikal "
              "antarkurva pada jumlah parameter yang setara merepresentasikan pengaruh murni "
              "rancangan arsitektur.")

    heading(doc, "1.6", "Metodologi Penelitian")
    body(doc, "Penelitian ini menggunakan pendekatan eksperimental kuantitatif. Korpus data "
              "disintesis secara terprogram dari kumpulan citra karakter Aksara Sunda terisolasi "
              "yang tersedia, dengan teks acuan diambil dari kosakata Sunda pada sumber data "
              "sekunder. Model yang diusulkan beserta seluruh model pembanding diimplementasikan "
              "dan dilatih di bawah protokol pelatihan yang identik, mencakup optimizer, "
              "penjadwal laju pembelajaran, jumlah epoch, skema augmentasi, dan pembagian data "
              "yang sama, sehingga perbedaan performa yang teramati dapat diatribusikan pada "
              "perbedaan rancangan model.")
    body(doc, "Setiap konfigurasi dijalankan pada sekurang-kurangnya tiga seed acak independen "
              "untuk memperoleh estimasi variabilitas, dan hasilnya dilaporkan sebagai rerata "
              "beserta simpangan bakunya. Pengujian signifikansi perbedaan antarmodel dilakukan "
              "menggunakan uji-t terkoreksi menurut Nadeau and Bengio (1999). Latensi diukur "
              "secara langsung pada perangkat keras nyata, bukan diestimasi dari jumlah operasi "
              "aritmetika, sesuai peringatan Ma et al. (2018).")

    heading(doc, "1.7", "Waktu dan Lokasi Penelitian")
    body(doc, "Penelitian dilaksanakan oleh penulis sebagai mahasiswa Program Studi S-1 Teknik "
              "Informatika, Departemen Ilmu Komputer, Fakultas Matematika dan Ilmu Pengetahuan "
              "Alam, Universitas Padjadjaran. Pelaksanaan eksperimen bertempat di laboratorium "
              "komputer program studi dengan memanfaatkan fasilitas komputasi yang tersedia.")
    body(doc, "Penelitian direncanakan berlangsung selama empat bulan terhitung sejak proposal "
              "ini disetujui, dengan rincian tahapan dan jadwal sebagaimana disajikan pada Bab IV.")


# =================================================================== BAB II
def bab_2(doc):
    bab(doc, "II", "TINJAUAN PUSTAKA")

    heading(doc, "2.1", "Aksara Sunda dan Upaya Pelestarian Digital", before=0)
    body(doc, "Aksara Sunda merupakan sistem tulisan turunan Brahmi yang digunakan untuk "
              "menuliskan bahasa Sunda. Sistem ini bersifat abugida, yaitu setiap aksara dasar "
              "melambangkan satu suku kata dengan vokal inheren yang dapat diubah melalui tanda "
              "diakritik. Bentuk aksara dasarnya terdiri atas aksara swara yang melambangkan "
              "vokal dan aksara ngalagena yang melambangkan konsonan bervokal inheren, ditambah "
              "sejumlah aksara untuk menuliskan bunyi serapan.")
    body(doc, "Adilazuarda et al. (2025) menempatkan Aksara Sunda sebagai salah satu dari "
              "delapan aksara Nusantara yang dihimpun dalam tolok ukur NusaAksara, bersama "
              "aksara Jawa, Bali, Batak, Lampung, Lontara, Jawi, dan Pegon. Penghimpunan tersebut "
              "dilakukan melalui anotasi pakar atas naskah terpindai, mencakup transkripsi ke "
              "aksara asli, transliterasi ke huruf Latin, dan terjemahan ke bahasa Indonesia. "
              "Ketersediaan lapisan transliterasi ini menjadikan NusaAksara sumber rujukan yang "
              "relevan bagi penelitian yang memerlukan teks acuan berbahasa Sunda.")
    body(doc, "Agustiansyah and Fauzi (2025) mencatat tiga tantangan khas pada pengenalan "
              "Aksara Sunda, yaitu variabilitas intrakelas yang tinggi akibat ragam gaya "
              "penulisan, kemiripan antarkelas yang tinggi pada sejumlah pasangan aksara, serta "
              "kelangkaan data yang persisten apabila dibandingkan dengan aksara modern seperti "
              "Latin.")

    heading(doc, "2.2", "Pengenalan Karakter Optik: Dari Terisolasi ke Sekuensial")
    body(doc, "Sistem pengenalan karakter optik pada awalnya dibangun secara bertahap, yaitu "
              "mendeteksi dan mengenali setiap karakter secara terpisah, kemudian merangkai "
              "hasilnya menjadi kata. Pendekatan tersebut menghadapi kendala mendasar yang "
              "dikenal sebagai paradoks Sayre, yaitu bahwa karakter pada tulisan bersambung tidak "
              "dapat disegmentasi tanpa terlebih dahulu dikenali, sementara pengenalannya "
              "memerlukan segmentasi. Shi et al. (2017) mencatat bahwa metode berbasis karakter "
              "menuntut pendeteksi karakter yang akurat, yang pada praktiknya boros sumber daya "
              "dan rentan galat akibat latar yang kompleks serta karakter yang saling bersentuhan.")
    body(doc, "Perkembangan berikutnya mengalihkan perumusan masalah dari klasifikasi menjadi "
              "pelabelan sekuens, sehingga segmentasi karakter tidak lagi diperlukan sebagai "
              "tahap terpisah. Peralihan inilah yang melandasi arsitektur pengenalan teks modern, "
              "dan menjelaskan mengapa tolok ukur berbasis karakter terisolasi relatif jarang "
              "menjadi fokus pengembangan sejak pertengahan dekade 2010-an.")
    body(doc, "Meskipun demikian, klasifikasi karakter terisolasi tetap relevan pada konteks "
              "aksara daerah dengan ketersediaan data terbatas, karena pengumpulan citra karakter "
              "tunggal jauh lebih terjangkau dibandingkan penganotasian teks utuh beserta "
              "transkripsinya. Kondisi inilah yang menjelaskan mengapa penelitian Aksara Sunda "
              "yang tersedia terkonsentrasi pada ranah terisolasi, sekaligus menjadi titik tolak "
              "penelitian ini untuk memanfaatkan korpus terisolasi tersebut sebagai bahan baku "
              "sintesis data sekuensial.")

    heading(doc, "2.3", "Convolutional Neural Network sebagai Ekstraktor Fitur")
    body(doc, "Convolutional Neural Network mengekstraksi fitur spasial melalui operasi "
              "konvolusi berbobot bersama, sehingga jumlah parameternya jauh lebih hemat "
              "dibandingkan lapisan terhubung penuh dengan daya representasi setara. Sejumlah "
              "teknik telah menjadi komponen baku pada pelatihannya, di antaranya normalisasi "
              "batch yang menstabilkan distribusi aktivasi antarlapisan (Ioffe and Szegedy, 2015) "
              "dan koneksi residual yang memungkinkan pelatihan jaringan berlapis banyak (He et "
              "al., 2016).")
    body(doc, "Pada ranah perangkat dengan sumber daya terbatas, berkembang keluarga arsitektur "
              "yang dirancang khusus untuk efisiensi. Iandola et al. (2016) menekan jumlah "
              "parameter melalui modul fire pada SqueezeNet, Sandler et al. (2018) memperkenalkan "
              "blok residual terbalik dengan bottleneck linear pada MobileNetV2, dan Howard et "
              "al. (2019) memadukan pencarian arsitektur otomatis dengan penyempurnaan manual "
              "pada MobileNetV3. Tan and Le (2019) merumuskan penskalaan majemuk yang menyelaraskan "
              "kedalaman, lebar, dan resolusi masukan secara proporsional pada EfficientNet.")
    body(doc, "Ma et al. (2018) memberikan koreksi metodologis penting terhadap keluarga "
              "arsitektur tersebut. Berdasarkan pengukuran langsung, mereka menunjukkan bahwa "
              "jumlah operasi aritmetika bukan proksi yang memadai bagi kecepatan sebenarnya, "
              "karena biaya akses memori dan tingkat paralelisme turut menentukan latensi. "
              "Temuan ini menjadi salah satu landasan penelitian ini untuk mengukur latensi "
              "secara langsung pada perangkat keras, bukan mengestimasinya.")

    heading(doc, "2.4", "Connectionist Temporal Classification")
    body(doc, "Connectionist Temporal Classification (CTC) yang diperkenalkan Graves et al. "
              "(2006) merupakan fungsi kerugian yang memungkinkan pelatihan model pelabelan "
              "sekuens tanpa memerlukan penyelarasan eksplisit antara masukan dan label. Model "
              "menghasilkan distribusi probabilitas atas himpunan label yang diperluas dengan "
              "satu simbol khusus bernama blank pada setiap langkah waktu.")
    body(doc, "Keluaran mentah tersebut dipetakan menjadi label akhir melalui aturan reduksi "
              "banyak-ke-satu, yaitu menggabungkan label identik yang berurutan kemudian "
              "menghapus seluruh simbol blank. Sebagai ilustrasi, urutan keluaran a-a-blank-b-b "
              "dan a-blank-b keduanya tereduksi menjadi label ab. Simbol blank berfungsi sebagai "
              "pemisah yang memungkinkan pengenalan label berulang, misalnya membedakan aa dari a.")
    body(doc, "Fungsi kerugian CTC didefinisikan sebagai negatif logaritma dari jumlah "
              "probabilitas atas seluruh penyelarasan yang tereduksi menjadi label acuan. Jumlah "
              "penyelarasan tersebut tumbuh secara eksponensial terhadap panjang sekuens, namun "
              "dapat dihitung secara efisien melalui algoritma maju-mundur dengan pemrograman "
              "dinamis. Sifat inilah yang menjadikan CTC dapat dilatih secara ujung-ke-ujung "
              "hanya dengan pasangan citra dan teks acuan, tanpa anotasi posisi per karakter.")
    caption(doc, "Gambar 2.1  Pemetaan banyak-ke-satu pada Connectionist Temporal Classification. "
                 "Beberapa penyelarasan keluaran per langkah waktu yang berbeda dapat tereduksi "
                 "menjadi label acuan yang sama.")
    body(doc, "Shi et al. (2017) memadukan CTC dengan ekstraktor fitur konvolusi dan lapisan "
              "rekuren dwiarah dalam arsitektur Convolutional Recurrent Neural Network (CRNN). "
              "Arsitektur ini memindai citra dari kiri ke kanan menjadi urutan vektor fitur, "
              "memodelkan kebergantungan antarlangkah melalui lapisan rekuren, kemudian "
              "menerapkan CTC pada keluarannya. Mereka melaporkan empat sifat yang menjadikannya "
              "sesuai untuk penerapan praktis, yaitu dapat dilatih ujung-ke-ujung, menangani "
              "sekuens dengan panjang bebas tanpa segmentasi karakter, tidak terikat pada "
              "leksikon tertentu, serta menghasilkan model yang jauh lebih ringkas.")

    heading(doc, "2.5", "Paradigma Dekoder Alternatif: Attention dan Transformer")
    body(doc, "Selain CTC, terdapat dua paradigma dekoder lain yang lazim digunakan pada "
              "pengenalan teks. Paradigma pertama adalah dekoder berbasis attention. Shi et al. "
              "(2016) mengusulkan arsitektur yang memadukan jaringan transformasi spasial untuk "
              "merektifikasi citra dengan jaringan pengenalan sekuens berbasis attention. "
              "Mekanisme attention memungkinkan dekoder memilih wilayah fitur yang relevan pada "
              "setiap langkah penerjemahan, sehingga secara implisit memodelkan kebergantungan "
              "antarkarakter.")
    body(doc, "Paradigma kedua adalah dekoder berbasis transformer. Li et al. (2023) "
              "mengusulkan TrOCR yang meniadakan tulang punggung konvolusi sepenuhnya dan "
              "menggunakan encoder bergaya Vision Transformer (Dosovitskiy et al., 2021) yang "
              "dipadukan dengan decoder bergaya model bahasa terlatih. Pendekatan ini melaporkan "
              "performa terdepan pada pengenalan teks cetak, tulisan tangan, dan teks pemandangan.")
    body(doc, "Perlu dicatat bahwa kedua paradigma alternatif tersebut umumnya bergantung pada "
              "ketersediaan data berskala besar maupun bobot terlatih dari korpus besar. Li et "
              "al. (2023) secara eksplisit memanfaatkan model terlatih untuk mengatasi kebutuhan "
              "tersebut. Pada ranah aksara daerah dengan data terbatas, ketergantungan ini "
              "menjadi pertimbangan risiko tersendiri yang diperhitungkan pada rancangan "
              "penelitian ini.")

    heading(doc, "2.6", "Sintesis Data untuk Pelatihan Model Pengenalan Teks")
    body(doc, "Sintesis data merupakan strategi baku pada ranah pengenalan teks untuk mengatasi "
              "keterbatasan data beranotasi. Jaderberg et al. (2014a) membangun mesin pembangkit "
              "teks sintetis yang menghasilkan data pelatihan dalam jumlah tak terbatas tanpa "
              "biaya penganotasian manusia, dan melatih pengenal kata sepenuhnya dari data "
              "tersebut. Gupta et al. (2016) memperluas pendekatan ini untuk tugas lokalisasi "
              "dengan menempelkan teks sintetis pada citra latar alami secara memperhatikan "
              "geometri adegan.")
    body(doc, "Kedua pekerjaan tersebut menegaskan bahwa data sintetis dapat memadai untuk "
              "melatih model pengenalan teks berkinerja tinggi, sepanjang proses sintesisnya "
              "menghadirkan keragaman yang cukup. Prinsip inilah yang diadopsi pada penelitian "
              "ini, dengan penyesuaian bahwa keragaman tidak diperoleh dari ragam fon "
              "komputasional, melainkan dari ragam instans citra karakter pada korpus terisolasi "
              "yang tersedia.")

    heading(doc, "2.7", "Efisiensi Model dan Optimalitas Pareto")
    body(doc, "Suatu konfigurasi model dikatakan optimal secara Pareto apabila tidak terdapat "
              "konfigurasi lain yang lebih unggul pada seluruh kriteria evaluasi secara "
              "bersamaan. Himpunan seluruh konfigurasi yang bersifat demikian membentuk frontier "
              "Pareto, yang merepresentasikan batas pertukaran terbaik yang dapat dicapai pada "
              "ruang kriteria tersebut.")
    caption(doc, "Gambar 2.2  Ilustrasi frontier Pareto pada ruang akurasi-latensi-ukuran model. "
                 "Titik yang tidak didominasi oleh titik lain pada ketiga sumbu membentuk "
                 "frontier.")
    body(doc, "Dehghani et al. (2022) mengkritisi praktik pengklaiman efisiensi yang bersandar "
              "pada satu metrik tunggal, karena jumlah parameter, jumlah operasi aritmetika, "
              "kecepatan, dan konsumsi memori kerap memberikan urutan peringkat yang saling "
              "bertentangan. Mereka menganjurkan pelaporan beberapa metrik secara bersamaan. "
              "Anjuran tersebut sejalan dengan temuan pengukuran Ma et al. (2018) dan menjadi "
              "dasar penetapan tiga sumbu evaluasi pada penelitian ini.")
    body(doc, "Pada ranah perbandingan arsitektur ringan, Shahriar (2026) melaksanakan studi "
              "terkendali lintas generasi yang membandingkan sembilan keluarga arsitektur di "
              "bawah protokol pelatihan dan evaluasi yang seragam, disertai analisis Pareto "
              "antara akurasi dan sumber daya. Studi tersebut menjadi preseden metodologis bagi "
              "rancangan eksperimen pada penelitian ini.")

    heading(doc, "2.8", "Penelitian Terdahulu yang Relevan")
    body(doc, "Agustiansyah and Fauzi (2025) membandingkan lima arsitektur pembelajaran "
              "mendalam pada klasifikasi 30 kelas Aksara Sunda terisolasi dan menguji "
              "ketahanannya terhadap data di luar sebaran. Mereka melaporkan pembalikan "
              "peringkat, yaitu EfficientNet-B0 unggul pada data sebaran dalam dengan akurasi "
              "96,9 persen namun menurun tajam pada data luar sebaran, sedangkan ResNet-50 yang "
              "lebih rendah pada sebaran dalam justru paling tahan dengan akurasi 92,5 persen "
              "pada data luar sebaran. Penelitian tersebut berbatas pada karakter terisolasi dan "
              "tidak mencakup pengenalan sekuensial maupun pengukuran latensi.")
    body(doc, "Adilazuarda et al. (2025) menghimpun tolok ukur NusaAksara dari 75 buku yang "
              "mencakup 7.137 halaman dan mengevaluasi beragam sistem pada tugas segmentasi, "
              "OCR, transliterasi, terjemahan, dan identifikasi bahasa. Untuk tugas OCR pada "
              "Aksara Sunda, seluruh sistem yang diuji menghasilkan Character Error Rate melebihi "
              "satu, termasuk PP-OCRv3 yang telah disetel khusus. Penelitian tersebut bersifat "
              "evaluatif dan tidak mengusulkan arsitektur baru maupun strategi penanggulangan "
              "keterbatasan data.")
    body(doc, "Pada ranah sistem OCR ringan secara umum, Zhang et al. (2026) melaporkan "
              "PP-OCRv6 dengan rentang 1,5 juta hingga 34,5 juta parameter yang mengungguli model "
              "bahasa visual berskala miliaran parameter pada tugas OCR. Sistem tersebut "
              "menunjukkan bahwa model ringan dapat bersaing pada ranah OCR, namun cakupan "
              "aksaranya tidak meliputi aksara Nusantara sebagaimana dibuktikan oleh hasil "
              "pengujian PP-OCR pada NusaAksara.")

    heading(doc, "2.9", "Kesenjangan Penelitian")
    body(doc, "Berdasarkan telaah pustaka di atas, teridentifikasi tiga kesenjangan. Pertama, "
              "penelitian pengenalan Aksara Sunda terkonsentrasi pada klasifikasi karakter "
              "terisolasi, sedangkan pengenalan pada tataran kata dan kalimat belum tersedia "
              "dalam bentuk yang layak pakai, sebagaimana dibuktikan oleh Character Error Rate "
              "melebihi satu pada seluruh sistem yang diuji Adilazuarda et al. (2025).")
    body(doc, "Kedua, strategi sintesis data yang telah terbukti efektif pada aksara Latin "
              "(Jaderberg et al., 2014a; Gupta et al., 2016) belum diterapkan untuk membangun "
              "korpus teks sekuensial Aksara Sunda, padahal keterbatasan data justru "
              "diidentifikasi sebagai penyebab utama kegagalan sistem terdahulu.")
    body(doc, "Ketiga, karakterisasi pertukaran antara akurasi, latensi terukur, dan ukuran "
              "model belum pernah dilakukan pada tugas pengenalan teks sekuensial Aksara Sunda, "
              "padahal karakterisasi tersebut menentukan kelayakan penerapan praktisnya. "
              "Penelitian ini diusulkan untuk menutup ketiga kesenjangan tersebut secara "
              "bersamaan.")


# ================================================================== BAB III
def bab_3(doc):
    bab(doc, "III", "METODE PENELITIAN")

    heading(doc, "3.1", "Asas Penelitian", before=0)
    body(doc, "Penelitian ini disusun dengan menjunjung asas dapat diulang (repeatable) dan "
              "dapat direproduksi (reproducible). Asas dapat diulang dipenuhi dengan menetapkan "
              "seed acak secara eksplisit pada setiap eksekusi, mencatat seluruh hiperparameter "
              "dalam berkas konfigurasi terversi, serta menyimpan bobot model dan berkas hasil "
              "setiap eksperimen. Asas dapat direproduksi dipenuhi dengan menggunakan sumber data "
              "yang tersedia untuk publik, mendokumentasikan prosedur sintesis data secara "
              "lengkap beserta kode pembangkitnya, dan melaporkan versi seluruh pustaka "
              "perangkat lunak yang digunakan.")
    body(doc, "Seluruh perbandingan antarmodel dilaksanakan di bawah protokol pelatihan dan "
              "evaluasi yang identik. Variabel yang tidak sedang diteliti, meliputi optimizer, "
              "penjadwal laju pembelajaran, jumlah epoch, skema augmentasi, dan pembagian data, "
              "dikendalikan konstan sehingga perbedaan performa yang teramati dapat diatribusikan "
              "pada perbedaan rancangan model.")

    heading(doc, "3.2", "Rancangan Penelitian")
    body(doc, "Penelitian dilaksanakan melalui enam tahap yang tersusun berurutan. Tahap "
              "pertama adalah penyiapan dan pembakuan korpus karakter Aksara Sunda terisolasi "
              "dari sumber-sumber yang tersedia. Tahap kedua adalah perancangan dan pelaksanaan "
              "prosedur sintesis citra teks sekuensial beserta labelnya. Tahap ketiga adalah "
              "implementasi model sekuensial berbasis CTC beserta model garis dasar. Tahap "
              "keempat adalah pelaksanaan rancangan eksperimen dua faktor. Tahap kelima adalah "
              "pengukuran latensi dan penyusunan frontier Pareto. Tahap keenam adalah analisis, "
              "pengujian signifikansi statistik, dan penarikan simpulan.")
    body(doc, "Pelaksanaan tahap ketiga hingga kelima disusun secara bertahap menurut tingkat "
              "risiko implementasi. Tahap inti mencakup model berbasis CTC beserta garis dasar "
              "segmentasi-lalu-klasifikasi, yang keduanya bersandar pada teknik yang telah mapan. "
              "Perluasan pada paradigma attention dan transformer dilaksanakan setelah tahap inti "
              "terselesaikan, dengan pertimbangan bahwa kedua paradigma tersebut memiliki risiko "
              "konvergensi yang lebih tinggi pada korpus berukuran sedang.")
    caption(doc, "Gambar 3.1  Diagram alir pelaksanaan penelitian. Penyelesaian tahap inti "
                 "berfungsi sebagai gerbang bagi pelaksanaan tahap perluasan.")

    heading(doc, "3.3", "Sumber dan Rancangan Data")
    heading(doc, "3.3.1", "Sumber Data", before=6, after=4)
    body(doc, "Data bersumber dari dua jenis sumber sekunder. Sumber pertama adalah korpus "
              "citra karakter Aksara Sunda terisolasi yang mencakup 30 kelas, meliputi aksara "
              "swara, aksara ngalagena, dan aksara untuk bunyi serapan. Korpus ini dihimpun dari "
              "repositori publik dan koleksi yang tersedia di lingkungan program studi, dengan "
              "tujuan memperoleh keragaman gaya penulisan yang memadai.")
    body(doc, "Sumber kedua adalah teks acuan berbahasa Sunda yang digunakan sebagai rangkaian "
              "target pada proses sintesis. Teks acuan diambil dari lapisan transliterasi Latin "
              "pada tolok ukur NusaAksara (Adilazuarda et al., 2025), yang telah melalui proses "
              "anotasi dan validasi pakar. Penggunaan teks acuan berbahasa Sunda yang sahih "
              "bertujuan memastikan bahwa rangkaian karakter yang disintesis merepresentasikan "
              "distribusi kemunculan karakter yang wajar, bukan permutasi acak yang tidak "
              "mencerminkan kaidah penulisan bahasa Sunda.")
    heading(doc, "3.3.2", "Prosedur Sintesis Citra Teks Sekuensial", before=6, after=4)
    body(doc, "Sintesis dilaksanakan melalui empat langkah. Langkah pertama adalah pemetaan "
              "setiap satuan tulisan pada teks acuan berhuruf Latin menjadi kelas aksara yang "
              "bersesuaian pada korpus terisolasi. Langkah kedua adalah pemilihan satu instans "
              "citra secara acak dari himpunan instans yang tersedia pada kelas tersebut. "
              "Pemilihan acak ini merupakan sumber utama keragaman, karena kata yang sama yang "
              "disintesis berulang kali akan menghasilkan citra yang berbeda-beda, sehingga model "
              "tidak dapat menghafal pola piksel tertentu.")
    body(doc, "Langkah ketiga adalah perangkaian instans-instans citra tersebut secara "
              "horizontal dengan penyeragaman tinggi citra dan penyisipan jarak antaraksara. "
              "Jarak antaraksara, pergeseran vertikal, serta variasi skala kecil diberikan secara "
              "acak dalam rentang terbatas untuk menghadirkan keragaman geometris. Langkah "
              "keempat adalah pencatatan label berupa urutan indeks kelas penyusun citra tersebut, "
              "yang bersifat eksak menurut konstruksinya sehingga tidak memerlukan penganotasian "
              "manual maupun verifikasi manusia.")
    caption(doc, "Gambar 3.2  Prosedur sintesis citra kata Aksara Sunda dari korpus karakter "
                 "terisolasi. Pemilihan instans secara acak pada setiap posisi menjadi sumber "
                 "keragaman visual.")
    body(doc, "Sintesis dilakukan pada dua tataran, yaitu tataran kata dan tataran frasa "
              "pendek. Pada tataran frasa, spasi antarkata diperlakukan sebagai satu kelas "
              "tersendiri sehingga model mempelajari batas antarkata sebagai bagian dari sekuens "
              "keluarannya.")
    heading(doc, "3.3.3", "Skema Pembagian Data", before=6, after=4)
    body(doc, "Pembagian data dirancang dengan dua lapis kendali untuk mencegah kebocoran "
              "informasi. Lapis pertama adalah pemisahan pada tataran kata, yaitu kata yang "
              "muncul pada himpunan uji dipastikan tidak pernah muncul pada himpunan latih. "
              "Kendali ini menguji kemampuan generalisasi model terhadap kombinasi karakter yang "
              "belum pernah ditemuinya, bukan sekadar kemampuan menghafal kata yang telah "
              "dipelajari.")
    body(doc, "Lapis kedua adalah pemisahan pada tataran instans citra karakter, yaitu instans "
              "citra yang digunakan untuk menyintesis data uji diambil dari subhimpunan yang "
              "terpisah dari instans yang digunakan pada data latih. Kombinasi kedua lapis "
              "kendali tersebut memastikan bahwa evaluasi dilakukan terhadap kata yang baru "
              "sekaligus terhadap perwujudan visual yang baru.")

    heading(doc, "3.4", "Arsitektur Model yang Diusulkan")
    heading(doc, "3.4.1", "Tulang Punggung Konvolusi", before=6, after=4)
    body(doc, "Tulang punggung yang diusulkan berupa jaringan konvolusi empat blok dengan lebar "
              "kanal 32, 64, 128, dan 128 pada konfigurasi dasar. Setiap blok tersusun atas "
              "konvolusi tiga kali tiga, normalisasi batch (Ioffe and Szegedy, 2015), dan "
              "aktivasi ReLU. Perbedaan mendasar terhadap tulang punggung untuk klasifikasi "
              "terletak pada skema penurunan resolusi, yaitu penurunan pada sumbu tinggi "
              "dilakukan hingga satu piksel sedangkan penurunan pada sumbu lebar dibatasi. Skema "
              "asimetris ini mempertahankan resolusi horizontal yang diperlukan sebagai sumbu "
              "waktu bagi pelabelan sekuens.")
    body(doc, "Peta fitur keluaran tulang punggung selanjutnya dibaca sebagai urutan vektor "
              "fitur, dengan setiap langkah waktu merepresentasikan satu irisan vertikal citra "
              "masukan. Spesifikasi lengkap lapisan beserta dimensi keluarannya disajikan pada "
              "Tabel 3.1, dengan W menyatakan lebar citra masukan yang bersifat variabel.")
    caption(doc, "Tabel 3.1  Spesifikasi lapisan arsitektur CRNN yang diusulkan", before=6, after=4)
    table(doc, [
        ["No", "Lapisan", "Konfigurasi", "Dimensi keluaran"],
        ["1", "Masukan", "citra grayscale", "32 x W x 1"],
        ["2", "Konv-1 + BN + ReLU", "3x3, 32 kanal, pad 1", "32 x W x 32"],
        ["3", "MaxPool-1", "2x2, langkah 2", "16 x W/2 x 32"],
        ["4", "Konv-2 + BN + ReLU", "3x3, 64 kanal, pad 1", "16 x W/2 x 64"],
        ["5", "MaxPool-2", "2x2, langkah 2", "8 x W/4 x 64"],
        ["6", "Konv-3 + BN + ReLU", "3x3, 128 kanal, pad 1", "8 x W/4 x 128"],
        ["7", "MaxPool-3", "2x1, langkah (2,1)", "4 x W/4 x 128"],
        ["8", "Konv-4 + BN + ReLU", "3x3, 128 kanal, pad 1", "4 x W/4 x 128"],
        ["9", "MaxPool-4", "2x1, langkah (2,1)", "2 x W/4 x 128"],
        ["10", "Pool adaptif tinggi", "tinggi menjadi 1", "1 x W/4 x 128"],
        ["11", "Pembentukan sekuens", "T = W/4 langkah", "T x 128"],
        ["12", "BiLSTM", "2 lapis, 256 unit tersembunyi", "T x 512"],
        ["13", "Proyeksi linear", "512 menjadi C+1", "T x (C+1)"],
        ["14", "Kerugian CTC", "C kelas + 1 blank", "-"],
    ], widths=[1.2, 3.6, 5.2, 4.0])
    heading(doc, "3.4.2", "Kepala Sekuensial Berbasis CTC", before=6, after=4)
    body(doc, "Urutan vektor fitur keluaran tulang punggung diteruskan ke lapisan rekuren "
              "dwiarah berjenis Long Short-Term Memory sebanyak dua lapis, yang berfungsi "
              "memodelkan kebergantungan antarlangkah waktu dari kedua arah. Keluaran lapisan "
              "rekuren diproyeksikan secara linear menjadi distribusi probabilitas atas C kelas "
              "aksara ditambah satu simbol blank pada setiap langkah waktu, kemudian dilatih "
              "menggunakan fungsi kerugian CTC (Graves et al., 2006).")
    body(doc, "Pada tahap inferensi, penerjemahan keluaran menjadi teks dilakukan melalui "
              "penerjemahan rakus (greedy decoding), yaitu pengambilan label berprobabilitas "
              "tertinggi pada setiap langkah waktu yang dilanjutkan dengan aturan reduksi CTC. "
              "Penerjemahan berbasis pencarian berkas (beam search) disertakan sebagai "
              "pembanding untuk menilai pertukaran antara akurasi dan latensi pada tahap "
              "penerjemahan.")
    heading(doc, "3.4.3", "Paradigma Dekoder Pembanding", before=6, after=4)
    body(doc, "Dua paradigma dekoder pembanding disertakan pada tahap perluasan. Paradigma "
              "pertama adalah dekoder berbasis attention yang mengacu pada rancangan Shi et al. "
              "(2016), dengan tulang punggung konvolusi yang sama sehingga perbedaan performa "
              "dapat diatribusikan pada mekanisme dekodernya. Paradigma kedua adalah dekoder "
              "berbasis transformer yang mengacu pada rancangan Li et al. (2023), dengan "
              "penyesuaian skala mengingat keterbatasan volume data.")

    heading(doc, "3.5", "Model Garis Dasar Segmentasi-lalu-Klasifikasi")
    body(doc, "Sebagai garis dasar disertakan pendekatan konvensional berupa segmentasi "
              "karakter yang diikuti klasifikasi terisolasi. Karena citra uji disintesis oleh "
              "penulis, posisi batas setiap karakter diketahui secara eksak, sehingga tahap "
              "segmentasi dapat dilaksanakan secara sempurna tanpa galat. Kondisi ini "
              "menempatkan garis dasar pada posisi yang menguntungkan, yaitu seolah-olah tersedia "
              "pendeteksi karakter yang sempurna.")
    body(doc, "Perbandingan terhadap garis dasar dalam kondisi tersebut memungkinkan penjawaban "
              "satu pertanyaan yang bersifat mendasar, yaitu seberapa besar keuntungan yang "
              "sesungguhnya diperoleh dari pemodelan sekuens apabila persoalan segmentasi "
              "diandaikan telah terselesaikan. Perbandingan ini juga memanfaatkan kembali model "
              "klasifikasi karakter terisolasi yang telah tersedia sehingga tidak memerlukan "
              "pelatihan tambahan yang berarti.")

    heading(doc, "3.6", "Rancangan Eksperimen Dua Faktor")
    body(doc, "Perbandingan antararsitektur pada ukuran bawaan masing-masing tidak dapat "
              "memisahkan pengaruh ukuran model dari pengaruh rancangan arsitektur, karena kedua "
              "variabel tersebut berubah secara bersamaan. Untuk mengatasi keterancuan tersebut, "
              "penelitian ini menerapkan rancangan eksperimen dua faktor.")
    body(doc, "Faktor pertama adalah rancangan arsitektur, yang mencakup topologi, jenis "
              "dekoder, dan skema penurunan resolusi. Faktor ini dikendalikan konstan di dalam "
              "satu kurva dan divariasikan antarkurva. Faktor kedua adalah pengali lebar kanal "
              "yang menskalakan jumlah kanal pada seluruh lapisan konvolusi secara proporsional, "
              "sebagaimana diterapkan Sandler et al. (2018) dan Tan and Le (2019). Konfigurasi "
              "pengali yang diteliti disajikan pada Tabel 3.2.")
    caption(doc, "Tabel 3.2  Konfigurasi pengali lebar kanal pada rancangan dua faktor",
            before=6, after=4)
    table(doc, [
        ["Pengali lebar", "Konv-1", "Konv-2", "Konv-3", "Konv-4"],
        ["0,25", "8", "16", "32", "32"],
        ["0,50", "16", "32", "64", "64"],
        ["0,75", "24", "48", "96", "96"],
        ["1,00 (dasar)", "32", "64", "128", "128"],
        ["1,50", "48", "96", "192", "192"],
    ], widths=[3.4, 2.6, 2.6, 2.6, 2.6])
    body(doc, "Dengan rancangan tersebut, setiap arsitektur menghasilkan satu kurva pada bidang "
              "jumlah parameter terhadap akurasi. Bentuk kurva merepresentasikan pengaruh murni "
              "ukuran model karena rancangan arsitektur dikendalikan konstan, sedangkan jarak "
              "vertikal antarkurva pada jumlah parameter yang setara merepresentasikan pengaruh "
              "murni rancangan arsitektur. Daftar seluruh model yang dibandingkan disajikan pada "
              "Tabel 3.3.")
    caption(doc, "Tabel 3.3  Model pembanding dan garis dasar", before=6, after=4)
    table(doc, [
        ["Kelompok", "Model", "Keterangan"],
        ["Diusulkan", "CRNN + CTC", "Lima konfigurasi pengali lebar kanal"],
        ["Perluasan", "Dekoder attention", "Tulang punggung identik, dekoder berbeda"],
        ["Perluasan", "Dekoder transformer", "Tulang punggung identik, dekoder berbeda"],
        ["Garis dasar", "Segmentasi + klasifikasi", "Segmentasi eksak dari konstruksi sintesis"],
        ["Pembanding", "Arsitektur ringan terpublikasi",
         "Sebagai tulang punggung, ukuran bawaan"],
    ], widths=[3.0, 4.6, 6.4])

    heading(doc, "3.7", "Prosedur Pelatihan dan Pengolahan Data Eksperimen")
    body(doc, "Pelatihan seluruh model menggunakan optimizer AdamW (Loshchilov and Hutter, "
              "2019) dengan penjadwal Cosine Annealing (Loshchilov and Hutter, 2017). Laju "
              "pembelajaran awal ditentukan melalui pencarian terbatas yang dijalankan dengan "
              "protokol identik untuk seluruh model, sehingga tidak ada model yang memperoleh "
              "keuntungan penyetelan yang tidak setara. Penghentian dini diterapkan berdasarkan "
              "Character Error Rate pada himpunan validasi.")
    body(doc, "Setiap konfigurasi dijalankan pada sekurang-kurangnya tiga seed acak independen. "
              "Hasil dilaporkan sebagai rerata beserta simpangan baku, dan perbandingan "
              "antarmodel diuji signifikansinya menggunakan uji-t terkoreksi menurut Nadeau and "
              "Bengio (1999), yang mengoreksi ketergantungan antarlipatan pada estimasi galat "
              "generalisasi. Pelaporan hasil dengan seed tunggal dihindari karena berpotensi "
              "mencerminkan keberuntungan inisialisasi, bukan keunggulan rancangan.")

    heading(doc, "3.8", "Pengukuran dan Kriteria Keberhasilan Penelitian")
    body(doc, "Performa pengenalan diukur menggunakan Character Error Rate sebagai metrik "
              "utama, yaitu jarak edit Levenshtein antara teks prediksi dan teks acuan yang "
              "dinormalisasi terhadap panjang teks acuan. Pemilihan metrik ini memungkinkan "
              "perbandingan langsung terhadap angka yang dilaporkan Adilazuarda et al. (2025). "
              "Metrik pendamping meliputi Word Error Rate dan akurasi kecocokan penuh pada "
              "tataran kata.")
    body(doc, "Efisiensi diukur melalui dua besaran. Besaran pertama adalah jumlah parameter "
              "model. Besaran kedua adalah latensi inferensi ujung-ke-ujung yang diukur secara "
              "langsung pada perangkat keras dengan ukuran batch satu, mencakup tahap "
              "prapemrosesan, propagasi maju, dan penerjemahan keluaran. Pengukuran dilakukan "
              "setelah tahap pemanasan dan diulang sejumlah kali untuk memperoleh nilai rerata "
              "beserta persentil ke-95.")
    caption(doc, "Tabel 3.4  Kriteria keberhasilan penelitian", before=6, after=4)
    table(doc, [
        ["No", "Kriteria", "Ambang"],
        ["1", "Korpus sintetis terbangun dan terdokumentasi",
         "Terverifikasi secara visual dan statistik"],
        ["2", "Character Error Rate model yang diusulkan",
         "Jauh di bawah 1,0 pada himpunan uji"],
        ["3", "Perbandingan terhadap garis dasar",
         "Terukur dan teruji signifikansinya"],
        ["4", "Kurva dua faktor untuk seluruh pengali lebar",
         "Lengkap pada minimal tiga seed"],
        ["5", "Frontier Pareto teridentifikasi",
         "Himpunan tak terdominasi pada tiga sumbu"],
    ], widths=[1.2, 7.0, 5.8])
    body(doc, "Penelitian dinyatakan berhasil apabila seluruh kriteria pada Tabel 3.4 "
              "terpenuhi. Perlu ditegaskan bahwa keberhasilan penelitian tidak disyaratkan pada "
              "tercapainya akurasi tertinggi oleh model yang diusulkan, melainkan pada "
              "terkarakterisasinya hubungan pertukaran antara ketiga sumbu evaluasi secara sahih "
              "dan dapat direproduksi.")

    heading(doc, "3.9", "Kebutuhan Perangkat Keras dan Perangkat Lunak")
    body(doc, "Kebutuhan perangkat keras meliputi unit pemroses grafis NVIDIA GeForce RTX 4070 "
              "berkapasitas memori 12 gigabyte untuk pelatihan model, serta unit pemroses pusat "
              "untuk pengukuran latensi pada kondisi yang merepresentasikan perangkat tanpa "
              "akselerator khusus. Ketersediaan beberapa unit komputasi memungkinkan pelaksanaan "
              "eksperimen secara paralel sehingga jadwal penelitian dapat dipenuhi.")
    body(doc, "Kebutuhan perangkat lunak meliputi pustaka PyTorch untuk implementasi dan "
              "pelatihan model beserta implementasi bawaan fungsi kerugian CTC, pustaka OpenCV "
              "dan Pillow untuk sintesis dan prapemrosesan citra, pustaka NumPy dan SciPy untuk "
              "analisis numerik dan pengujian statistik, serta sistem kendali versi Git untuk "
              "menjamin keterlacakan kode dan konfigurasi eksperimen.")

    heading(doc, "3.10", "Batasan Penelitian")
    body(doc, "Penelitian ini memiliki tiga batasan yang perlu dinyatakan secara terbuka. "
              "Pertama, data pelatihan dan pengujian berupa citra teks hasil sintesis dari "
              "korpus karakter terisolasi, bukan citra naskah terpindai maupun foto teks pada "
              "kondisi nyata. Dengan demikian, hasil penelitian ini berlaku untuk teks Aksara "
              "Sunda tercetak hasil sintesis, dan perluasannya pada naskah nyata memerlukan "
              "penelitian lanjutan.")
    body(doc, "Kedua, penelitian ini berbatas pada tahap pengenalan dan tidak mencakup tahap "
              "deteksi wilayah teks, karena citra masukan telah tersedia dalam keadaan terpotong "
              "pada tataran kata menurut konstruksi sintesisnya. Ketiga, konfigurasi yang "
              "teridentifikasi optimal secara Pareto bersifat spesifik terhadap korpus dan "
              "karakteristik data pada penelitian ini; penerapannya pada distribusi data yang "
              "berbeda memerlukan pengulangan prosedur yang sama, bukan penggunaan ulang model "
              "secara langsung.")


# =================================================================== BAB IV
def bab_4(doc):
    bab(doc, "IV", "JADWAL PENELITIAN")

    heading(doc, "4.1", "Tahapan Pelaksanaan Penelitian", before=0)
    body(doc, "Pelaksanaan penelitian dibagi ke dalam delapan kegiatan pokok yang terentang "
              "selama empat bulan. Kegiatan pertama adalah studi pustaka lanjutan dan pembakuan "
              "korpus karakter Aksara Sunda terisolasi dari seluruh sumber yang tersedia. "
              "Kegiatan kedua adalah perancangan, implementasi, dan verifikasi prosedur sintesis "
              "citra teks sekuensial beserta korpus yang dihasilkannya.")
    body(doc, "Kegiatan ketiga adalah implementasi model sekuensial berbasis CTC beserta "
              "verifikasi kebenaran pipeline pelatihan pada subhimpunan data berukuran kecil. "
              "Kegiatan keempat adalah implementasi model garis dasar segmentasi-lalu-klasifikasi. "
              "Kegiatan kelima adalah pelaksanaan eksperimen dua faktor pada seluruh konfigurasi "
              "pengali lebar kanal beserta pengukuran latensinya.")
    body(doc, "Kegiatan keenam adalah pelaksanaan tahap perluasan berupa implementasi dan "
              "pengujian dekoder attention dan transformer, yang dijadwalkan setelah tahap inti "
              "terselesaikan. Kegiatan ketujuh adalah analisis hasil, penyusunan frontier Pareto, "
              "dan pengujian signifikansi statistik. Kegiatan kedelapan adalah penyusunan laporan "
              "tugas akhir dan persiapan sidang.")

    heading(doc, "4.2", "Jadwal Kegiatan")
    body(doc, "Rincian jadwal pelaksanaan setiap kegiatan disajikan pada Tabel 4.1. Tanda "
              "silang menyatakan bulan pelaksanaan kegiatan yang bersangkutan.")
    caption(doc, "Tabel 4.1  Jadwal pelaksanaan penelitian", before=6, after=4)
    table(doc, [
        ["No", "Kegiatan", "Bulan 1", "Bulan 2", "Bulan 3", "Bulan 4"],
        ["1", "Studi pustaka dan pembakuan korpus karakter", "X", "", "", ""],
        ["2", "Perancangan dan pelaksanaan sintesis data", "X", "X", "", ""],
        ["3", "Implementasi model sekuensial berbasis CTC", "", "X", "", ""],
        ["4", "Implementasi model garis dasar", "", "X", "", ""],
        ["5", "Eksperimen dua faktor dan pengukuran latensi", "", "X", "X", ""],
        ["6", "Tahap perluasan: attention dan transformer", "", "", "X", ""],
        ["7", "Analisis, frontier Pareto, dan uji statistik", "", "", "X", "X"],
        ["8", "Penyusunan laporan dan persiapan sidang", "", "", "", "X"],
    ], widths=[1.0, 7.0, 1.5, 1.5, 1.5, 1.5], font_pt=10)


# =========================================================== DAFTAR PUSTAKA
REFERENSI = [
    "Adilazuarda, M.F., Wijanarko, M.I., Susanto, L., Nur'aini, K., Wijaya, D. and Aji, A.F. "
    "(2025) 'NusaAksara: a multimodal and multilingual benchmark for preserving Indonesian "
    "indigenous scripts', in Proceedings of the 63rd Annual Meeting of the Association for "
    "Computational Linguistics (ACL 2025). Association for Computational Linguistics.",

    "Agustiansyah, Y. and Fauzi, D.R. (2025) 'From local features to global context: comparing "
    "CNN and transformer for Sundanese script classification', Journal of Intelligent Systems "
    "Technology and Informatics, 1(2), pp. 53-61.",

    "Bergstra, J. and Bengio, Y. (2012) 'Random search for hyper-parameter optimization', "
    "Journal of Machine Learning Research, 13, pp. 281-305.",

    "Dehghani, M., Arnab, A., Beyer, L., Vaswani, A. and Tay, Y. (2022) 'The efficiency "
    "misnomer', in International Conference on Learning Representations (ICLR 2022).",

    "Dosovitskiy, A., Beyer, L., Kolesnikov, A., Weissenborn, D., Zhai, X., Unterthiner, T., "
    "Dehghani, M., Minderer, M., Heigold, G., Gelly, S., Uszkoreit, J. and Houlsby, N. (2021) "
    "'An image is worth 16x16 words: transformers for image recognition at scale', in "
    "International Conference on Learning Representations (ICLR 2021).",

    "Graves, A., Fernandez, S., Gomez, F. and Schmidhuber, J. (2006) 'Connectionist temporal "
    "classification: labelling unsegmented sequence data with recurrent neural networks', in "
    "Proceedings of the 23rd International Conference on Machine Learning (ICML). New York: ACM, "
    "pp. 369-376.",

    "Gupta, A., Vedaldi, A. and Zisserman, A. (2016) 'Synthetic data for text localisation in "
    "natural images', in Proceedings of the IEEE Conference on Computer Vision and Pattern "
    "Recognition (CVPR), pp. 2315-2324.",

    "He, K., Zhang, X., Ren, S. and Sun, J. (2016) 'Deep residual learning for image "
    "recognition', in Proceedings of the IEEE Conference on Computer Vision and Pattern "
    "Recognition (CVPR), pp. 770-778.",

    "Howard, A., Sandler, M., Chu, G., Chen, L.-C., Chen, B., Tan, M., Wang, W., Zhu, Y., Pang, "
    "R., Vasudevan, V., Le, Q.V. and Adam, H. (2019) 'Searching for MobileNetV3', in Proceedings "
    "of the IEEE/CVF International Conference on Computer Vision (ICCV), pp. 1314-1324.",

    "Iandola, F.N., Han, S., Moskewicz, M.W., Ashraf, K., Dally, W.J. and Keutzer, K. (2016) "
    "'SqueezeNet: AlexNet-level accuracy with 50x fewer parameters and <0.5MB model size', arXiv "
    "preprint arXiv:1602.07360.",

    "Ioffe, S. and Szegedy, C. (2015) 'Batch normalization: accelerating deep network training "
    "by reducing internal covariate shift', in Proceedings of the 32nd International Conference "
    "on Machine Learning (ICML), pp. 448-456.",

    "Jaderberg, M., Simonyan, K., Vedaldi, A. and Zisserman, A. (2014a) 'Synthetic data and "
    "artificial neural networks for natural scene text recognition', arXiv preprint "
    "arXiv:1406.2227.",

    "Jaderberg, M., Vedaldi, A. and Zisserman, A. (2014b) 'Deep features for text spotting', in "
    "European Conference on Computer Vision (ECCV). Cham: Springer, pp. 512-528.",

    "Li, M., Lv, T., Chen, J., Cui, L., Lu, Y., Florencio, D., Zhang, C., Li, Z. and Wei, F. "
    "(2023) 'TrOCR: transformer-based optical character recognition with pre-trained models', in "
    "Proceedings of the AAAI Conference on Artificial Intelligence, 37(11), pp. 13094-13102.",

    "Loshchilov, I. and Hutter, F. (2017) 'SGDR: stochastic gradient descent with warm restarts', "
    "in International Conference on Learning Representations (ICLR 2017).",

    "Loshchilov, I. and Hutter, F. (2019) 'Decoupled weight decay regularization', in "
    "International Conference on Learning Representations (ICLR 2019).",

    "Ma, N., Zhang, X., Zheng, H.-T. and Sun, J. (2018) 'ShuffleNet V2: practical guidelines for "
    "efficient CNN architecture design', in European Conference on Computer Vision (ECCV). Cham: "
    "Springer, pp. 122-138.",

    "Nadeau, C. and Bengio, Y. (1999) 'Inference for the generalization error', in Advances in "
    "Neural Information Processing Systems 12 (NIPS 1999), pp. 307-313.",

    "Sandler, M., Howard, A., Zhu, M., Zhmoginov, A. and Chen, L.-C. (2018) 'MobileNetV2: "
    "inverted residuals and linear bottlenecks', in Proceedings of the IEEE Conference on "
    "Computer Vision and Pattern Recognition (CVPR), pp. 4510-4520.",

    "Shahriar, T. (2026) 'Do newer lightweight CNNs perform better under resource constraints? A "
    "controlled multigenerational study of architecture, initialization, training budget, and "
    "efficiency', arXiv preprint arXiv:2607.01984.",

    "Shi, B., Bai, X. and Yao, C. (2017) 'An end-to-end trainable neural network for image-based "
    "sequence recognition and its application to scene text recognition', IEEE Transactions on "
    "Pattern Analysis and Machine Intelligence, 39(11), pp. 2298-2304.",

    "Shi, B., Wang, X., Lyu, P., Yao, C. and Bai, X. (2016) 'Robust scene text recognition with "
    "automatic rectification', in Proceedings of the IEEE Conference on Computer Vision and "
    "Pattern Recognition (CVPR), pp. 4168-4176.",

    "Tan, M. and Le, Q.V. (2019) 'EfficientNet: rethinking model scaling for convolutional "
    "neural networks', in Proceedings of the 36th International Conference on Machine Learning "
    "(ICML), pp. 6105-6114.",

    "Zhang, Y., Wang, X., Lin, M., Zhang, Y. and Deng, P. (2026) 'PP-OCRv6: from 1.5M to 34.5M "
    "parameters, surpassing billion-scale VLMs on OCR tasks', arXiv preprint arXiv:2606.13108.",
]


def daftar_pustaka(doc):
    doc.add_page_break()
    center(doc, "DAFTAR PUSTAKA", bold=True, spacing=1.5, after=12)
    for r in REFERENSI:
        pustaka(doc, r)


def riwayat_hidup(doc):
    doc.add_page_break()
    center(doc, "RIWAYAT HIDUP", bold=True, spacing=1.5, after=12)
    body(doc, "[Bagian ini diisi oleh penulis: nama lengkap, tempat dan tanggal lahir, "
              "riwayat pendidikan formal, pengalaman organisasi dan kepanitiaan, prestasi "
              "akademik maupun nonakademik, serta pengalaman kerja atau magang apabila ada. "
              "Riwayat hidup bersifat wajib pada skripsi Program Studi Teknik Informatika.]",
         indent=False)


# ====================================================================== main
def main():
    doc = Document()
    setup(doc)
    halaman_judul(doc)
    kata_pengantar(doc)
    daftar(doc, "DAFTAR ISI", DAFTAR_ISI)
    daftar(doc, "DAFTAR GAMBAR", DAFTAR_GAMBAR)
    daftar(doc, "DAFTAR TABEL", DAFTAR_TABEL)
    bab_1(doc)
    bab_2(doc)
    bab_3(doc)
    bab_4(doc)
    daftar_pustaka(doc)
    riwayat_hidup(doc)
    doc.save(OUT)
    print("SAVED:", OUT)
    print("paragraf:", len(doc.paragraphs), "tabel:", len(doc.tables))


if __name__ == "__main__":
    main()
